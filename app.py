import threading
import time
from flask import Flask, redirect, url_for, render_template, request, send_file
from docx import Document
from datetime import datetime
import io, os
from dateutil import parser
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from flask_dance.contrib.google import make_google_blueprint, google
from flask_login import LoginManager, login_user, login_required, current_user, UserMixin
from docx2pdf import convert


app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "default-secret-key")
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'


# Google OAuth setup
google_bp = make_google_blueprint(
    client_id=os.environ.get("GOOGLE_OAUTH_CLIENT_ID"),
    client_secret=os.environ.get("GOOGLE_OAUTH_CLIENT_SECRET"),
    redirect_url="/login/google/authorized",
    scope=["https://www.googleapis.com/auth/userinfo.email","https://www.googleapis.com/auth/userinfo.profile","openid"]
)
app.register_blueprint(google_bp, url_prefix="/login")

# User class for session management
class User(UserMixin):
    def __init__(self, email):
        self.id = email
        self.email = email

login_manager = LoginManager(app)

@login_manager.user_loader
def load_user(user_id):
    return User(user_id)

# OAuth + login flow
@app.route("/")
def index():
    if not google.authorized:
        return redirect(url_for("google.login"))

    resp = google.get("/oauth2/v2/userinfo")
    email = resp.json()["email"]

    # Optional domain check
    #allowed_domains = ["myschool.edu"]
    #if not any(email.endswith(f"@{domain}") for domain in allowed_domains):
     #   return "Unauthorized access", 403

    user = User(email)
    login_user(user)
    return redirect(url_for("syllabus_form"))

@app.route("/login/google/authorized")
def google_authorized():
    if not google.authorized:
        return redirect(url_for("index"))
    
    resp = google.get("/oauth2/v2/userinfo")
    email = resp.json()["email"]

    user = User(email)
    login_user(user)
    return redirect(url_for("syllabus_form"))

# Protected route for your syllabus form
@app.route("/generate-document", methods=["GET", "POST"])
@login_required
def syllabus_form():
    if request.method == "POST":
        form_data = request.form.to_dict()
        form_data["submitted_by"] = current_user.email
        # Proceed with document generation
        #return f"Form submitted by {form_data['submitted_by']}"
    


# @app.route('/generate-document', methods=['POST'])
# def generate_document():
        form = request.form
        class_name = form.get('className')
        assessment = form.get('assessmentTitle').upper()

        subject_map = {
            "LKG": ["Urdu", "English", "Mathematics"],
            "UKG": ["Urdu", "English", "Mathematics"],
            "First": ["Urdu", "English", "Mathematics"],
            "Second": ["Urdu", "English", "Mathematics", "Telugu"],
            "Third": ["Urdu", "English", "Mathematics", "Telugu"],
            "Fourth": ["Urdu", "English", "Mathematics", "Telugu", "EVS"],
            "Fifth": ["Urdu", "English", "Mathematics", "Telugu", "EVS"]
        }

        subjects = subject_map.get(class_name, [])

        # Extract dates from form
        dates = []
        syllabus_data = []
        for subject in subjects:
            key = subject.lower()
            date_str = form.get(f'{key}Date')
            try:
                date_obj = parser.parse(date_str)
                dates.append(date_obj)
                day = date_obj.strftime('%A')
            except Exception:
                date_obj = None
                day = form.get(f'{key}Day', '')
            syllabus_data.append({
                'date': date_str or '',
                'day': day,
                'subject': subject,
                'portion': form.get(f'{key}Syllabus', '')
            })

        latest_date = max(dates) if dates else None
        month = latest_date.strftime('%B') if latest_date else 'Month'
        year = latest_date.strftime('%Y') if latest_date else 'Year'

        #print(f"Generating document for {class_name} - {assessment} for {month} {year}")

        doc = Document("syllabus-template.docx")

        # 🔁 Replace placeholders in paragraphs while preserving formatting
        replacements = {
            "{assessment}": assessment,
            "{class}": class_name,
            "(month)": month,
            "(year)": year
        }

        for para in doc.paragraphs:
            for run in para.runs:
                #print(f"Processing paragraph: {run.text}")
                for key, val in replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, val)

        # 📋 Fill the first table (assuming headers already exist)
        table = doc.tables[0]
        start_row = 1  # Assuming first row is header


        # Sort based on parsed date field
        sorted_data = sorted(syllabus_data, key=lambda x: datetime.strptime(x['date'], '%Y-%m-%d'))

        for i, item in enumerate(sorted_data):
            if i + start_row < len(table.rows):
                row = table.rows[i + start_row]
            else:
                row = table.add_row()

            values = [item['date'], item['day'], item['subject'], item['portion']]

            for col_idx, val in enumerate(values):
                cell = row.cells[col_idx]
                cell.text = ""  # Clear any auto text

                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(val)
                # 🔁 Match header style (except portion column)
                header_run = table.rows[0].cells[col_idx].paragraphs[0].runs[0]
                #run.font.size = Pt(12)
                #run.font.name = 'Calibri'
                run.font.name = header_run.font.name
                run.font.size = Pt(16)
                run.bold = False if col_idx == 3 else True

                # 📐 Styling rules
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                if col_idx == 3:
                    # Portion column: left-aligned + wrap text
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        # Stable output folder
        output_dir = os.path.join(app.root_path, "static", "generated_docs")
        os.makedirs(output_dir, exist_ok=True)

        # Construct file paths
        docx_path = os.path.join(output_dir, f"{class_name}_Syllabus.docx")
        pdf_path = os.path.join(output_dir, f"{class_name}_Syllabus.pdf")

        # ⚙️ Save DOCX
        doc.save(docx_path)
        time.sleep(1)  # Let the OS catch up

        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"DOCX not saved properly to {docx_path}")

        # 📄 Convert to PDF
        convert(docx_path, pdf_path)

        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF not generated at {pdf_path}")

        # 💾 Serve PDF and schedule cleanup
        with open(pdf_path, "rb") as pdf_file:
            response = send_file(
                io.BytesIO(pdf_file.read()),
                as_attachment=True,
                download_name=f"{class_name}_Syllabus.pdf",
                mimetype="application/pdf"
            )

        delete_file_delayed(docx_path)
        delete_file_delayed(pdf_path)

        return response

        # buffer = io.BytesIO()
        # doc.save(buffer)
        # buffer.seek(0)

        # filename = f"{class_name}_Syllabus.docx"
        # return send_file(buffer, as_attachment=True, download_name=filename)
    return render_template("form.html")

def delete_file_delayed(path, delay=5):
    def delayed_delete():
        time.sleep(delay)
        if os.path.exists(path):
            os.remove(path)
    threading.Thread(target=delayed_delete).start()

if __name__ == '__main__':
    app.run(debug=True)
